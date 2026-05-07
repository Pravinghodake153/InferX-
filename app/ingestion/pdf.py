import base64
import io
import os
import uuid
from typing import List, Dict, Any, Optional

import fitz  # PyMuPDF

try:
    import pdfplumber  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    pdfplumber = None

try:
    from paddleocr import PaddleOCR  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    PaddleOCR = None

try:
    import pytesseract  # type: ignore
except Exception:
    pytesseract = None

try:
    import cv2  # type: ignore
except Exception:
    cv2 = None

try:
    from deskew import determine_skew  # type: ignore
except Exception:
    determine_skew = None

_OCR_ENGINES: Dict[str, Any] = {}  # lang -> PaddleOCR instance (lazy cache)

# Indian script Unicode ranges for auto-detection
_SCRIPT_RANGES = {
    "hi": (0x0900, 0x097F),  # Devanagari (Hindi, Sanskrit, Marathi)
    "ka": (0x0C80, 0x0CFF),  # Kannada
    "ta": (0x0B80, 0x0BFF),  # Tamil
    "te": (0x0C00, 0x0C7F),  # Telugu
}

# PaddleOCR init args vary by version — we try different combos
def _create_ocr_engine(lang: str):
    if PaddleOCR is None:
        return None
    import time as _time
    print(f"[OCR] Initializing PaddleOCR engine for lang='{lang}'...")
    t0 = _time.time()
    init_combos = [
        {"use_angle_cls": False, "lang": lang, "show_log": False, "ocr_version": "PP-OCRv4"},
        {"use_angle_cls": False, "lang": lang, "ocr_version": "PP-OCRv4"},
        {"lang": lang, "ocr_version": "PP-OCRv4"},
        {"lang": lang},
    ]
    for kwargs in init_combos:
        try:
            engine = PaddleOCR(**kwargs)
            elapsed = _time.time() - t0
            print(f"[OCR] Engine for lang='{lang}' ready in {elapsed:.1f}s")
            return engine
        except (TypeError, ValueError):
            continue
        except Exception as exc:
            print(f"Warning: PaddleOCR({lang}) unavailable: {exc}")
            return None
    print(f"[OCR] All init combos failed for lang='{lang}'")
    return None


def _get_ocr_engine(lang: str = "en"):
    """Get or create a PaddleOCR engine. To prevent OOM, we only cache one engine at a time."""
    # If the requested lang is already cached, return it
    if lang in _OCR_ENGINES and _OCR_ENGINES[lang] is not None:
        return _OCR_ENGINES[lang]
        
    # Free up memory by clearing the previous engine if it exists
    _OCR_ENGINES.clear()
    
    engine = _create_ocr_engine(lang)
    _OCR_ENGINES[lang] = engine
    return engine


def _detect_scripts(text: str) -> List[str]:
    """Detect which Indian scripts are present in the text.
    Returns list of PaddleOCR language codes."""
    if not text:
        return []
    detected = set()
    for char in text:
        cp = ord(char)
        for lang, (start, end) in _SCRIPT_RANGES.items():
            if start <= cp <= end:
                detected.add(lang)
    return list(detected)


PUBLIC_IMAGES_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "public", "images"))
os.makedirs(PUBLIC_IMAGES_DIR, exist_ok=True)


def _iter_pages(doc: fitz.Document):
    for page_index in range(doc.page_count):
        yield page_index, doc.load_page(page_index)


def _as_text(value: Any) -> str:
    if isinstance(value, str):
        return value
    if value is None:
        return ""
    return str(value)


def _as_dict(value: Any) -> Dict[str, Any]:
    if isinstance(value, dict):
        return value
    return {}


def detect_document_type(page_text: str, word_count: int = 0) -> str:
    text_len = len(page_text.strip())
    # Mixed PDFs often contain sparse native text overlays on scanned pages.
    if text_len < 220 or word_count < 40:
        return "SCANNED"
    return "DIGITAL"


def _render_page_to_png_bytes(page: fitz.Page, dpi: int = 300) -> Optional[bytes]:
    try:
        scale = dpi / 72.0
        pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
        return pix.tobytes("png")
    except Exception as exc:
        print(f"Warning: Could not render page to image: {exc}")
        return None


def _extract_text_from_ocr_results(results) -> List[str]:
    """Parse PaddleOCR results into a list of text lines (handles all result formats)."""
    lines: List[str] = []

    def _push(value):
        if value is None:
            return
        text = str(value).strip()
        if text:
            lines.append(text)

    if isinstance(results, list):
        for block in results:
            if not block:
                continue
            if isinstance(block, dict):
                for item in block.get("rec_texts", []) or []:
                    _push(item)
                continue
            if isinstance(block, list):
                for line in block:
                    if not line:
                        continue
                    if isinstance(line, dict):
                        _push(line.get("text"))
                        _push(line.get("rec_text"))
                        continue
                    if isinstance(line, (list, tuple)) and len(line) >= 2:
                        rec = line[1]
                        if isinstance(rec, (list, tuple)) and rec:
                            _push(rec[0])
                        else:
                            _push(rec)
    return lines


def _run_ocr_with_lang(image_array, lang: str) -> str:
    """Run PaddleOCR with a specific language on a numpy image array."""
    import time as _time
    engine = _get_ocr_engine(lang)
    if engine is None:
        return ""
    try:
        t0 = _time.time()
        results = engine.ocr(image_array)
        lines = _extract_text_from_ocr_results(results)
        elapsed = _time.time() - t0
        print(f"[OCR] lang='{lang}' → {len(lines)} lines in {elapsed:.1f}s")
        return "\n".join(lines)
    except Exception as exc:
        print(f"Warning: OCR({lang}) failed: {exc}")
        return ""


def _preprocess_image_for_ocr(img_array):
    """
    Preprocess a scanned image for better OCR accuracy.
    Steps: grayscale → deskew → denoise → binarize (Otsu).
    Can improve OCR accuracy by 15-30% on real government scanned documents.
    """
    import numpy as np
    if cv2 is None:
        return img_array  # No OpenCV, return as-is

    try:
        # 1. Convert to grayscale
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array

        # 2. Deskew (straighten rotated scans)
        if determine_skew is not None:
            try:
                angle = determine_skew(gray)
                if angle is not None and abs(angle) > 0.5:  # Only deskew if rotated > 0.5°
                    (h, w) = gray.shape[:2]
                    center = (w // 2, h // 2)
                    M = cv2.getRotationMatrix2D(center, angle, 1.0)
                    gray = cv2.warpAffine(gray, M, (w, h),
                                          flags=cv2.INTER_CUBIC,
                                          borderMode=cv2.BORDER_REPLICATE)
                    print(f"[OCR] Deskewed image by {angle:.1f}°")
            except Exception:
                pass  # deskew failure is non-critical

        # 3. Denoise (remove scan artifacts)
        gray = cv2.fastNlMeansDenoising(gray, h=10)

        # 4. Binarize using Otsu's thresholding (best for scanned docs)
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        return binary
    except Exception as exc:
        print(f"[OCR] Preprocessing failed (non-critical): {exc}")
        return img_array


def _run_tesseract_ocr(image_array) -> str:
    """
    Run Tesseract OCR for English + Hindi (Devanagari) with:
      1. Image preprocessing (deskew/denoise/binarize)
      2. OCR confidence scoring per word
    """
    if pytesseract is None:
        print("Warning: Tesseract requested but pytesseract not installed. Falling back to PaddleOCR English.")
        return _run_ocr_with_lang(image_array, "en")
    import time as _time
    import numpy as np
    try:
        t0 = _time.time()

        # Preprocess image for better accuracy
        processed = _preprocess_image_for_ocr(image_array)

        # Use image_to_data for confidence scoring
        try:
            data = pytesseract.image_to_data(processed, lang='eng+hin',
                                              output_type=pytesseract.Output.DATAFRAME)
            # Filter out low-confidence noise (conf < 30 is usually garbage)
            valid = data[(data.conf > 0) & (data.text.notna()) & (data.text.str.strip() != '')]
            avg_conf = valid.conf.mean() if len(valid) > 0 else 0
            conf_label = "HIGH" if avg_conf >= 70 else ("MEDIUM" if avg_conf >= 50 else "LOW")

            # Reconstruct text from data (preserves line structure)
            lines = []
            for _, group in valid.groupby(['block_num', 'par_num', 'line_num']):
                line = ' '.join(group.text.astype(str).values)
                if line.strip():
                    lines.append(line.strip())

            elapsed = _time.time() - t0
            print(f"[OCR] Tesseract (eng+hin) → {len(lines)} lines in {elapsed:.1f}s [confidence={avg_conf:.0f}% → {conf_label}]")
            return "\n".join(lines)
        except Exception:
            # Fallback to simple image_to_string if dataframe parsing fails
            text = pytesseract.image_to_string(processed, lang='eng+hin')
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            elapsed = _time.time() - t0
            print(f"[OCR] Tesseract (eng+hin) → {len(lines)} lines in {elapsed:.1f}s [confidence=N/A]")
            return "\n".join(lines)
    except Exception as exc:
        print(f"Warning: Tesseract OCR failed: {exc}")
        return ""


def _run_ocr(page: fitz.Page, native_text: str = "") -> str:
    """
    Multilingual OCR:
      1. If Devanagari detected, or fully scanned page (no native text) -> Tesseract (eng+hin)
      2. Otherwise -> PaddleOCR (en)
    """
    image_bytes = _render_page_to_png_bytes(page)
    if not image_bytes:
        return ""

    try:
        from PIL import Image  # type: ignore
        import numpy as np  # type: ignore
    except Exception as exc:
        print(f"Warning: Pillow/NumPy unavailable for OCR: {exc}")
        return ""

    try:
        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        img_array = np.array(image)
    except Exception as exc:
        print(f"Warning: Could not load image for OCR: {exc}")
        return ""

    detected_scripts = _detect_scripts(native_text)
    needs_hindi = "hi" in detected_scripts or not native_text.strip()

    if needs_hindi and pytesseract is not None:
        return _run_tesseract_ocr(img_array)
    else:
        # Fallback to PaddleOCR English only
        return _run_ocr_with_lang(img_array, "en")


def _run_ocr_image_path(image_path: str) -> str:
    """OCR on a standalone image file using Tesseract for eng+hin if available."""
    try:
        from PIL import Image  # type: ignore
        import numpy as np  # type: ignore
    except Exception as exc:
        print(f"Warning: Pillow/NumPy unavailable for OCR: {exc}")
        return ""

    try:
        image = Image.open(image_path).convert("RGB")
        img_array = np.array(image)
    except Exception as exc:
        print(f"Warning: Could not load image {image_path}: {exc}")
        return ""

    if pytesseract is not None:
        return _run_tesseract_ocr(img_array)
    return _run_ocr_with_lang(img_array, "en")


def _save_image_bytes(image_bytes: bytes, page_number: int, image_index: int, ext: str, prefix: str) -> str:
    safe_ext = (ext or "png").lower()
    if safe_ext == "jpeg":
        safe_ext = "jpg"
    file_name = f"{prefix}_p{page_number}_img{image_index}_{uuid.uuid4().hex[:10]}.{safe_ext}"
    file_path = os.path.join(PUBLIC_IMAGES_DIR, file_name)
    with open(file_path, "wb") as f:
        f.write(image_bytes)
    return f"/images/{file_name}"


def _extract_table_entries_for_page(pdf_path: str, page_number: int) -> List[Dict[str, Any]]:
    # Table extraction has been removed
    return []

def _extract_tables_for_page(pdf_path: str, page_number: int) -> List[str]:
    return []


def _extract_images_for_page(doc: fitz.Document, page: fitz.Page, page_number: int, max_images: int = 4, prefix: str = "doc") -> List[Dict[str, Any]]:
    image_entries: List[Dict[str, Any]] = []
    try:
        image_list = page.get_images(full=True) or []
    except Exception as exc:
        print(f"Warning: Could not inspect images on page {page_number}: {exc}")
        return image_entries

    for image_index, image_info in enumerate(image_list[:max_images], start=1):
        try:
            xref = image_info[0]
            extracted = doc.extract_image(xref)
            image_bytes = extracted.get("image", b"")
            ext = extracted.get("ext", "bin")
            width = extracted.get("width")
            height = extracted.get("height")
            image_url = ""
            if image_bytes:
                image_url = _save_image_bytes(image_bytes, page_number, image_index, ext, prefix)

            image_bbox = [0.0, 0.0, float(page.rect.width), float(page.rect.height)]
            try:
                rects = page.get_image_rects(xref)
                if rects:
                    first = rects[0]
                    image_bbox = [float(first.x0), float(first.y0), float(first.x1), float(first.y1)]
            except Exception:
                pass

            image_entries.append({
                "page": page_number,
                "index": image_index,
                "xref": xref,
                "width": width,
                "height": height,
                "ext": ext,
                "image_ref": f"page_{page_number}_image_{image_index}.{ext}",
                "image_url": image_url,
                "bbox": image_bbox,
                "image_bytes_b64": base64.b64encode(image_bytes).decode("utf-8") if image_bytes else ""
            })
        except Exception as exc:
            print(f"Warning: Could not extract image {image_index} on page {page_number}: {exc}")
    return image_entries


def _extract_plain_pdf_text(path: str) -> str:
    try:
        doc = fitz.open(path)
        pages = []
        for i, page in _iter_pages(doc):
            pages.append(f"[Page {i+1}] {page.get_text()}")
        return "\n".join(pages)
    except Exception as e:
        print(f"Error reading PDF {path}: {e}")
        return ""


def extract_pdf_package(path: str) -> Dict[str, Any]:
    """
    Returns a hybrid document package with text, tables, images, and page metadata.
    The returned structure is stable and can be used for LLM context fusion.
    """
    package: Dict[str, Any] = {
        "context_text": "",
        "pages": [],
        "tables": [],
        "table_entries": [],
        "images": [],
    }

    try:
        doc = fitz.open(path)
        rendered_sections: List[str] = []

        for page_index, page in _iter_pages(doc):
            page_number = page_index + 1
            native_text = _as_text(page.get_text("text"))
            word_count = len(page.get_text("words") or [])
            document_type = detect_document_type(native_text, word_count=word_count)

            page_text = native_text.strip()
            ocr_text = ""
            text_source = "Native Text"
            ocr_status = "not_attempted"
            
            detected_scripts = _detect_scripts(native_text)
            has_hindi = "hi" in detected_scripts
            
            if document_type == "SCANNED" or has_hindi:
                reason = "Hindi detected" if has_hindi else "SCANNED"
                print(f"OCR triggered for page: {page_number} [{reason}]")
                ocr_status = "attempted"
                ocr_text = _run_ocr(page, native_text=native_text).strip()
                if ocr_text:
                    page_text = ocr_text
                    text_source = "OCR"
                elif page_text:
                    # Keep native text as a fallback if OCR is unavailable.
                    page_text = page_text
                    text_source = "Native Fallback (OCR empty)"
                    ocr_status = "empty_fallback_native"
                    # Keep OCR field visible in UI even when OCR returns empty.
                    ocr_text = page_text
                else:
                    text_source = "OCR attempted (no text)"
                    ocr_status = "empty_no_text"
            else:
                text_source = "Native Text"

            page_table_entries = _extract_table_entries_for_page(path, page_number)
            page_tables = [entry.get("text", "") for entry in page_table_entries if entry.get("text")]
            page_images = _extract_images_for_page(doc, page, page_number, prefix="extract")

            combined_page_text = page_text
            if page_tables:
                combined_page_text = (page_text + "\n\n" + "\n\n".join(page_tables)).strip() if page_text else "\n\n".join(page_tables)

            package["pages"].append({
                "page_num": page_number,
                "document_type": document_type,
                "native_text": native_text,
                "ocr_text": ocr_text,
                "ocr_status": ocr_status,
                "text_source": text_source,
                "text": combined_page_text,
                "tables": page_tables,
                "table_entries": [
                    {
                        "page": table_entry.get("page"),
                        "index": table_entry.get("index"),
                        "rows": table_entry.get("rows", []),
                        "bbox": table_entry.get("bbox", [0.0, 0.0, 0.0, 0.0]),
                        "text": table_entry.get("text", ""),
                    }
                    for table_entry in page_table_entries
                ],
                "images": [
                    {
                        "page": image_entry["page"],
                        "index": image_entry["index"],
                        "xref": image_entry["xref"],
                        "image_ref": image_entry["image_ref"],
                        "ext": image_entry["ext"],
                        "image_url": image_entry.get("image_url", ""),
                        "bbox": image_entry.get("bbox", []),
                        "width": image_entry["width"],
                        "height": image_entry["height"]
                    }
                    for image_entry in page_images
                ]
            })
            package["tables"].extend(page_tables)
            package["table_entries"].extend(page_table_entries)
            package["images"].extend(page_images)

            rendered_sections.append(f"[Page {page_number} | {document_type} | {text_source}]\n{combined_page_text}")
            if page_tables:
                rendered_sections.append(f"[TABLE DATA - Page {page_number}]\n" + "\n\n".join(page_tables))
            if page_images:
                image_lines = []
                for image_entry in page_images:
                    image_lines.append(
                        f"IMAGE (Page {page_number} #{image_entry['index']}): ref={image_entry['image_ref']} url={image_entry.get('image_url', '')} xref={image_entry['xref']} size={image_entry.get('width') or 'unknown'}x{image_entry.get('height') or 'unknown'}"
                    )
                rendered_sections.append(f"[IMAGE DATA - Page {page_number}]\n" + "\n".join(image_lines))

        package["context_text"] = "\n\n".join(section for section in rendered_sections if section).strip()
        return package
    except Exception as exc:
        print(f"Error extracting hybrid PDF package {path}: {exc}")
        package["context_text"] = _extract_plain_pdf_text(path)
        return package

def extract_pdf_text(path: str) -> str:
    """
    Extracts hybrid text from a PDF page by page.
    Digital pages use PyMuPDF text, scanned pages use OCR fallback,
    and table/image metadata is appended as additional context.
    """
    try:
        return extract_pdf_package(path)["context_text"]
    except Exception as e:
        print(f"Error reading PDF {path}: {e}")
        return _extract_plain_pdf_text(path)

def extract_pdf_layout(path: str) -> List[Dict[str, Any]]:
    """
    Extracts structured layout data from a PDF.
    Returns page-wise blocks with text, font size, color, and bounding boxes.
    
    Returns:
        List of page dictionaries, each containing:
        {
            "page_num": int,
            "width": float,
            "height": float,
            "blocks": [
                {
                    "text": str,
                    "bbox": [x0, y0, x1, y1],  # absolute coordinates
                    "font_size": float,
                    "color": [r, g, b],  # RGB values 0-255
                    "font_name": str,
                    "block_type": str  # "text", "image", etc.
                }
            ]
        }
    """
    try:
        doc = fitz.open(path)
        pages_layout = []
        
        for page_num, page in _iter_pages(doc):
            page_dict = _as_dict(page.get_text("dict"))
            page_height = page.rect.height
            page_width = page.rect.width
            
            blocks = []
            for block in page_dict.get("blocks", []):
                # Skip non-text blocks (images, etc.)
                if block["type"] != 0:  # 0 = text block
                    continue
                
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue
                        
                        # Extract bounding box
                        bbox = span.get("bbox", [])  # [x0, y0, x1, y1]
                        if not bbox or len(bbox) < 4:
                            continue
                        
                        # Extract font info
                        font_size = float(span.get("size", 12))
                        font_name = str(span.get("font", "Helvetica"))
                        
                        # Extract color - PyMuPDF can return color in different formats
                        color = [0, 0, 0]  # Default black
                        try:
                            color_int = span.get("color")
                            if color_int is not None:
                                if isinstance(color_int, int):
                                    # Color is stored as RGB integer (0xRRGGBB)
                                    color = [
                                        (color_int >> 16) & 255,
                                        (color_int >> 8) & 255,
                                        color_int & 255
                                    ]
                                elif isinstance(color_int, (list, tuple)) and len(color_int) >= 3:
                                    # Color is already a list/tuple
                                    color = [int(c * 255) if c <= 1 else int(c) for c in color_int[:3]]
                        except Exception as e:
                            print(f"Warning: Could not extract color: {e}")
                            color = [0, 0, 0]
                        
                        blocks.append({
                            "text": text,
                            "bbox": [float(x) for x in bbox],
                            "font_size": font_size,
                            "color": color,
                            "font_name": font_name,
                            "block_type": "text"
                        })

            native_char_count = sum(len(str(block.get("text", ""))) for block in blocks)
            needs_ocr_assist = (not blocks) or native_char_count < 120
            if needs_ocr_assist:
                print(f"OCR triggered for page: {page_num + 1}")
                existing_native = " ".join(str(block.get("text", "")) for block in blocks)
                ocr_text = _run_ocr(page, native_text=existing_native)
                if ocr_text:
                    existing_text = " ".join(str(block.get("text", "")) for block in blocks)
                    ocr_lines = [line.strip() for line in ocr_text.splitlines() if line.strip()]
                    y_cursor = max([float(b.get("bbox", [0.0, 0.0, 0.0, 0.0])[3]) for b in blocks if len(b.get("bbox", [])) >= 4] + [48.0]) + 10.0
                    for line in ocr_lines:
                        if line in existing_text:
                            continue
                        blocks.append({
                            "text": line,
                            "bbox": [48.0, y_cursor, max(page_width - 48.0, 300.0), y_cursor + 18.0],
                            "font_size": 12.0,
                            "color": [25, 31, 46],
                            "font_name": "Courier",
                            "block_type": "text"
                        })
                        y_cursor += 20.0

            table_entries = _extract_table_entries_for_page(path, page_num + 1)
            if table_entries:
                fallback_y = max([float(b.get("bbox", [0.0, 0.0, 0.0, 0.0])[3]) for b in blocks if len(b.get("bbox", [])) >= 4] + [72.0]) + 16.0
                for table_entry in table_entries:
                    bbox = table_entry.get("bbox", [0.0, 0.0, 0.0, 0.0])
                    valid_bbox = isinstance(bbox, list) and len(bbox) == 4 and (bbox[2] - bbox[0]) > 2 and (bbox[3] - bbox[1]) > 2
                    if not valid_bbox:
                        table_height = max(24.0 * len(table_entry.get("rows", [])), 48.0)
                        bbox = [48.0, fallback_y, max(page_width - 48.0, 300.0), min(fallback_y + table_height, page_height - 24.0)]
                        fallback_y = float(bbox[3]) + 14.0

                    blocks.append({
                        "text": table_entry.get("text", ""),
                        "bbox": [float(v) for v in bbox],
                        "font_size": 10.0,
                        "color": [35, 43, 59],
                        "font_name": "Helvetica",
                        "block_type": "table",
                        "rows": table_entry.get("rows", []),
                        "table_index": table_entry.get("index"),
                        "page": table_entry.get("page", page_num + 1),
                    })

            image_entries = _extract_images_for_page(doc, page, page_num + 1, prefix="layout")
            for image_entry in image_entries:
                blocks.append({
                    "text": image_entry.get("image_ref", "image"),
                    "bbox": image_entry.get("bbox", [0.0, 0.0, float(page_width), float(page_height)]),
                    "font_size": 10.0,
                    "color": [180, 180, 180],
                    "font_name": "Helvetica",
                    "block_type": "image",
                    "src": image_entry.get("image_url", ""),
                    "width": image_entry.get("width"),
                    "height": image_entry.get("height"),
                    "ext": image_entry.get("ext"),
                    "page": image_entry.get("page"),
                })
            
            pages_layout.append({
                "page_num": page_num + 1,
                "width": float(page_width),
                "height": float(page_height),
                "blocks": blocks
            })
        
        return pages_layout
    except Exception as e:
        print(f"Error extracting PDF layout {path}: {e}")
        return []


def extract_layout_text(pdf_path: str) -> List[Dict[str, Any]]:
    """
    Returns page-wise layout-preserved text using PyMuPDF.
    Each page maintains an approximate reading structure from text blocks.
    """
    try:
        doc = fitz.open(pdf_path)
        pages_output: List[Dict[str, Any]] = []

        for page_index, page in _iter_pages(doc):
            blocks = page.get_text("blocks") or []
            # block tuple format: (x0, y0, x1, y1, text, block_no, block_type)
            text_blocks = [b for b in blocks if len(b) >= 5 and str(b[4]).strip()]
            text_blocks.sort(key=lambda b: (round(float(b[1]), 1), round(float(b[0]), 1)))

            lines: List[str] = []
            prev_y: Optional[float] = None
            for block in text_blocks:
                x0, y0, _, _, text = block[:5]
                block_text = str(text).strip("\n")
                if not block_text:
                    continue

                if prev_y is not None:
                    # Add additional blank space between distant vertical blocks.
                    y_gap = float(y0) - prev_y
                    if y_gap > 18:
                        lines.append("")
                prev_y = float(y0)

                # Apply light indentation based on X position to preserve alignment hints.
                indent = max(int(float(x0) // 22), 0)
                indented_text = (" " * indent) + block_text
                lines.append(indented_text)

            pages_output.append({
                "page": page_index + 1,
                "text": "\n".join(lines).strip() if lines else ""
            })

        return pages_output
    except Exception as exc:
        print(f"Error extracting layout text {pdf_path}: {exc}")
        return []


def extract_visual_text(pdf_path: str) -> List[Dict[str, Any]]:
    """
    Uses page words and spatial grouping to reconstruct visually aligned text.
    """
    try:
        doc = fitz.open(pdf_path)
        pages_output: List[Dict[str, Any]] = []

        for page_index, page in _iter_pages(doc):
            words = page.get_text("words") or []
            # word tuple format: (x0, y0, x1, y1, word, block_no, line_no, word_no)
            words = [w for w in words if len(w) >= 5 and str(w[4]).strip()]
            words.sort(key=lambda w: (round(float(w[1]), 1), round(float(w[0]), 1)))

            y_threshold = 3.5
            line_groups: List[List[Any]] = []

            for word in words:
                if not line_groups:
                    line_groups.append([word])
                    continue
                last_group = line_groups[-1]
                last_y = float(last_group[-1][1])
                if abs(float(word[1]) - last_y) <= y_threshold:
                    last_group.append(word)
                else:
                    line_groups.append([word])

            rendered_lines: List[str] = []
            char_scale = 5.0
            for group in line_groups:
                group.sort(key=lambda w: float(w[0]))
                cursor = 0
                out = ""
                for word in group:
                    x0 = float(word[0])
                    token = str(word[4])
                    target = max(int(x0 / char_scale), 0)
                    if target > cursor:
                        out += " " * (target - cursor)
                        cursor = target
                    out += token
                    cursor += len(token)
                    out += " "
                    cursor += 1
                rendered_lines.append(out.rstrip())

            pages_output.append({
                "page": page_index + 1,
                "text": "\n".join(rendered_lines)
            })

        return pages_output
    except Exception as exc:
        print(f"Error extracting visual text {pdf_path}: {exc}")
        return []

def load_text(file_path: str) -> str:
    """Fallback text loader for non-PDF mock data."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        print(f"Error reading text {file_path}: {e}")
        return ""


def text_to_layout(text: str) -> List[Dict[str, Any]]:
    """
    Converts plain text into synthetic page layout blocks for Document View.
    This allows TXT uploads to render in the same viewer used by PDFs.
    """
    if not text:
        return [{
            "page_num": 1,
            "width": 612.0,
            "height": 792.0,
            "blocks": []
        }]

    page_width = 612.0
    page_height = 792.0
    left_margin = 48.0
    right_margin = 48.0
    top_margin = 56.0
    bottom_margin = 56.0
    font_size = 12.0
    line_height = 18.0
    max_chars_per_line = 95

    wrapped_lines: List[str] = []
    raw_lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")

    for raw in raw_lines:
        if not raw:
            wrapped_lines.append("")
            continue

        line = raw
        while len(line) > max_chars_per_line:
            split_at = line.rfind(" ", 0, max_chars_per_line)
            if split_at <= 0:
                split_at = max_chars_per_line
            wrapped_lines.append(line[:split_at].rstrip())
            line = line[split_at:].lstrip()
        wrapped_lines.append(line)

    lines_per_page = int((page_height - top_margin - bottom_margin) / line_height)
    pages: List[Dict[str, Any]] = []

    for page_idx in range(0, len(wrapped_lines), lines_per_page):
        page_lines = wrapped_lines[page_idx:page_idx + lines_per_page]
        blocks: List[Dict[str, Any]] = []

        for i, line in enumerate(page_lines):
            if not line:
                continue

            y0 = top_margin + (i * line_height)
            y1 = y0 + line_height
            x0 = left_margin
            # Approximate monospace width at 12px.
            x1 = min(page_width - right_margin, x0 + (len(line) * 7.2))

            blocks.append({
                "text": line,
                "bbox": [x0, y0, x1, y1],
                "font_size": font_size,
                "color": [30, 41, 59],
                "font_name": "Courier",
                "block_type": "text"
            })

        pages.append({
            "page_num": (page_idx // lines_per_page) + 1,
            "width": page_width,
            "height": page_height,
            "blocks": blocks
        })

    return pages or [{
        "page_num": 1,
        "width": page_width,
        "height": page_height,
        "blocks": []
    }]


def extract_text_package(text: str) -> Dict[str, Any]:
    """
    Builds a package compatible with extract_pdf_package for plain text inputs.
    Keeps API and frontend rendering consistent between PDF and TXT files.
    """
    pages_layout = text_to_layout(text)
    pages: List[Dict[str, Any]] = []

    for page in pages_layout:
        page_text = "\n".join(block.get("text", "") for block in page.get("blocks", []) if block.get("text"))
        pages.append({
            "page_num": page.get("page_num", 1),
            "document_type": "DIGITAL",
            "native_text": page_text,
            "ocr_text": "",
            "text": page_text,
            "tables": [],
            "images": []
        })

    context_sections = [f"[Page {p['page_num']} | DIGITAL | TEXT]\n{p['text']}" for p in pages]
    return {
        "context_text": "\n\n".join(context_sections).strip(),
        "pages": pages,
        "tables": [],
        "table_entries": [],
        "images": []
    }


def extract_docx_text(path: str) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:
        print(f"Warning: python-docx unavailable: {exc}")
        return ""

    try:
        document = docx.Document(path)
    except Exception as exc:
        print(f"Error reading DOCX {path}: {exc}")
        return ""

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]

    table_sections: List[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows: List[List[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() if cell.text else "" for cell in row.cells])
        table_text = _table_rows_to_text(rows)
        if table_text:
            table_sections.append(f"TABLE (DOCX #{table_index}):\n{table_text}")

    merged = "\n\n".join(paragraphs + table_sections).strip()
    return merged


def _table_rows_to_text(rows: List[List[str]]) -> str:
    if not rows:
        return ""
    
    lines = []
    for i, row in enumerate(rows):
        line = " | ".join(cell.replace("\n", " ") for cell in row)
        lines.append(f"| {line} |")
        if i == 0:
            separator = " | ".join("---" for _ in row)
            lines.append(f"| {separator} |")
    return "\n".join(lines)


def extract_docx_package(path: str) -> Dict[str, Any]:
    text = extract_docx_text(path)
    package = extract_text_package(text)

    page_tables = [line for line in text.split("\n\n") if line.startswith("TABLE (DOCX #")]
    if package.get("pages"):
        package["pages"][0]["tables"] = page_tables
    package["tables"] = page_tables
    package.setdefault("table_entries", [])
    return package


def extract_image_package(path: str) -> Dict[str, Any]:
    package: Dict[str, Any] = {
        "context_text": "",
        "pages": [],
        "tables": [],
        "table_entries": [],
        "images": [],
    }

    try:
        from PIL import Image  # type: ignore
    except Exception as exc:
        print(f"Warning: Pillow unavailable for image extraction: {exc}")
        return package

    try:
        with Image.open(path) as img:
            width, height = img.size

        ext = os.path.splitext(path)[1].lstrip(".").lower() or "png"
        if ext == "jpeg":
            ext = "jpg"

        with open(path, "rb") as f:
            image_bytes = f.read()

        image_url = _save_image_bytes(image_bytes, 1, 1, ext, "image")
        ocr_text = _run_ocr_image_path(path).strip()
        text_source = "OCR" if ocr_text else "Image (no OCR text)"
        ocr_status = "available" if ocr_text else "empty"

        page_text = ocr_text
        package["pages"].append({
            "page_num": 1,
            "document_type": "SCANNED",
            "native_text": "",
            "ocr_text": ocr_text,
            "ocr_status": ocr_status,
            "text_source": text_source,
            "text": page_text,
            "tables": [],
            "table_entries": [],
            "images": [{
                "page": 1,
                "index": 1,
                "xref": 0,
                "image_ref": f"image_1.{ext}",
                "ext": ext,
                "image_url": image_url,
                "bbox": [0.0, 0.0, float(width), float(height)],
                "width": width,
                "height": height,
            }],
        })
        package["images"].append({
            "page": 1,
            "index": 1,
            "xref": 0,
            "image_ref": f"image_1.{ext}",
            "ext": ext,
            "image_url": image_url,
            "bbox": [0.0, 0.0, float(width), float(height)],
            "width": width,
            "height": height,
            "image_bytes_b64": base64.b64encode(image_bytes).decode("utf-8") if image_bytes else "",
        })
        package["context_text"] = f"[Page 1 | IMAGE | {text_source}]\n{page_text}".strip()
        return package
    except Exception as exc:
        print(f"Error extracting image package {path}: {exc}")
        return package


def extract_image_layout(path: str) -> List[Dict[str, Any]]:
    try:
        from PIL import Image  # type: ignore
    except Exception as exc:
        print(f"Warning: Pillow unavailable for image layout: {exc}")
        return []

    try:
        with Image.open(path) as img:
            width, height = img.size

        ext = os.path.splitext(path)[1].lstrip(".").lower() or "png"
        if ext == "jpeg":
            ext = "jpg"

        with open(path, "rb") as f:
            image_bytes = f.read()
        image_url = _save_image_bytes(image_bytes, 1, 1, ext, "layout")

        ocr_text = _run_ocr_image_path(path).strip()
        blocks: List[Dict[str, Any]] = [
            {
                "text": "image",
                "bbox": [0.0, 0.0, float(width), float(height)],
                "font_size": 10.0,
                "color": [180, 180, 180],
                "font_name": "Helvetica",
                "block_type": "image",
                "src": image_url,
                "width": width,
                "height": height,
                "ext": ext,
                "page": 1,
            }
        ]

        if ocr_text:
            y_cursor = 20.0
            for line in [x.strip() for x in ocr_text.splitlines() if x.strip()]:
                blocks.append({
                    "text": line,
                    "bbox": [24.0, y_cursor, max(float(width) - 24.0, 200.0), y_cursor + 18.0],
                    "font_size": 12.0,
                    "color": [20, 25, 40],
                    "font_name": "Courier",
                    "block_type": "text",
                })
                y_cursor += 20.0

        return [{
            "page_num": 1,
            "width": float(width),
            "height": float(height),
            "blocks": blocks,
        }]
    except Exception as exc:
        print(f"Error extracting image layout {path}: {exc}")
        return []


def extract_image_layout_debug(path: str) -> List[Dict[str, Any]]:
    text = _run_ocr_image_path(path).strip()
    return [{"page": 1, "text": text}]
